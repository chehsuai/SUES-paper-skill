#!/usr/bin/env python3
"""
学术论文 Word 文档生成器 — Academic Paper Word Generator

功能：
  - 生成符合国内学术规范的 Word 文档（.docx）
  - 正文宋体小四，1.5倍行距，首行缩进2字符
  - 三级标题（一、/（一）/ 1.）顶格，不加粗
  - 引用标记上标显示，Ctrl+点击跳转到对应参考文献
  - 参考文献格式符合 GB/T 7714

依赖：pip install python-docx lxml

用法示例见文件末尾。
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree
import re

# ── XML 命名空间 ──────────────────────────────────────────────
WML_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
XML_NS  = 'http://www.w3.org/XML/1998/namespace'


def _tag(tagname: str) -> str:
    """生成带 wml 命名空间的 XML 标签名。"""
    return f'{{{WML_NS}}}{tagname}'


# ── 书签 ID 计数器 ────────────────────────────────────────────
_bookmark_id = 0


def _next_bm_id() -> int:
    global _bookmark_id
    _bookmark_id += 1
    return _bookmark_id


def _reset_bm_id():
    global _bookmark_id
    _bookmark_id = 0


# ═══════════════════════════════════════════════════════════════
# 核心类
# ═══════════════════════════════════════════════════════════════

class PaperGenerator:
    """
    学术论文 Word 文档生成器。

    使用方式：
        gen = PaperGenerator()
        gen.set_cover_info(course='...', school='...', student_id='...',
                           author='...', advisor='...')
        gen.set_title('主标题', subtitle='副标题（可选）')
        gen.set_abstract('摘要内容...', keywords='关键词1；关键词2')
        gen.add_heading('一、绪论')
        gen.add_heading('（一）研究背景', level=2)
        gen.add_body('正文段落，引用用 [1] [2] 格式...')
        gen.add_reference('[1] 作者.题名[J].刊名, 年份, 卷(期): 页码.')
        gen.save('output.docx')
    """

    def __init__(self):
        self.doc = Document()
        self._setup_styles()
        self._refs_added = False
        _reset_bm_id()

    # ── 样式初始化 ──────────────────────────────────────────
    def _setup_styles(self):
        style = self.doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)  # 小四
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        pf = style.paragraph_format
        pf.line_spacing = 1.5
        pf.first_line_indent = Cm(0.74)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        for sec in self.doc.sections:
            sec.top_margin = Cm(2.54)
            sec.bottom_margin = Cm(2.54)
            sec.left_margin = Cm(3.18)
            sec.right_margin = Cm(3.18)

    # ── 基础工具 ────────────────────────────────────────────
    def _add_run(self, para, text, font_name='宋体', font_size=Pt(12),
                 bold=False, superscript=False):
        run = para.add_run(text)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = font_size
        run.bold = bold
        if superscript:
            run.font.superscript = True
        return run

    def _blank_line(self):
        p = self.doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(0)
        return p

    # ═══════════════════════════════════════════════════════
    # 公开 API
    # ═══════════════════════════════════════════════════════

    # ── 封面 ──────────────────────────────────────────────
    def set_cover_info(self, course: str, school: str,
                       student_id: str, author: str, advisor: str):
        """设置封面信息（在调用 set_title 之前调用）。"""
        for _ in range(6):
            self._blank_line()
        # 标题位置由 set_title 填入，这里留空
        self._cover_course = course
        self._cover_school = school
        self._cover_sid = student_id
        self._cover_author = author
        self._cover_advisor = advisor

    def set_title(self, title: str, subtitle: str = ''):
        """设置论文主标题（黑体二号加粗居中），可选副标题。"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(22)
        run.bold = True

        if subtitle:
            p2 = self.doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run(subtitle)
            run2.font.name = '黑体'
            run2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run2.font.size = Pt(22)
            run2.bold = True
        else:
            # 避免空白段落
            pass

        for _ in range(4):
            self._blank_line()

        # 封面信息
        for line in [
            f'课程名称：{getattr(self, "_cover_course", "（请填写）")}',
            f'学    院：{getattr(self, "_cover_school", "（请填写）")}',
            f'学    号：{getattr(self, "_cover_sid", "xxx")}',
            f'姓    名：{getattr(self, "_cover_author", "xxx")}',
            f'指导教师：{getattr(self, "_cover_advisor", "xxx")}'
        ]:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(line)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(14)
            run.bold = True

        self.doc.add_page_break()

    def set_inner_header(self, title: str, author_line: str):
        """
        内页标题行（黑体三号加粗居中）+ 作者行（宋体五号居中）。

        调用后自动进入正文区域。
        """
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(title)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(16)
        run.bold = True

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(author_line)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10.5)

    def set_abstract(self, abstract_text: str, keywords: str):
        """
        添加摘要和关键词。

        Parameters
        ----------
        abstract_text : 摘要正文（不含"摘要："前缀）
        keywords : 关键词字符串，用分号分隔，如 "品牌传播；沉浸式体验"
        """
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.first_line_indent = Cm(0.74)
        para.paragraph_format.space_after = Pt(4)
        r1 = para.add_run('摘要：')
        r1.font.name = '宋体'
        r1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r1.font.size = Pt(12)
        r1.bold = True
        r2 = para.add_run(abstract_text)
        r2.font.name = '宋体'
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r2.font.size = Pt(12)

        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.first_line_indent = Cm(0.74)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(8)
        r1 = para.add_run('关键词：')
        r1.font.name = '宋体'
        r1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r1.font.size = Pt(12)
        r1.bold = True
        r2 = para.add_run(keywords)
        r2.font.name = '宋体'
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r2.font.size = Pt(12)

    # ── 标题 ──────────────────────────────────────────────
    def add_heading(self, text: str, level: int = 1):
        """
        添加正文标题。

        level 1: 一、二、三...
        level 2: （一）（二）（三）...
        level 3: 1. 2. 3...
        所有标题：顶格，宋体小四，不加粗。
        """
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(3)
        run = para.add_run(text)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)
        run.bold = False
        return para

    # ── 正文段落（含引用处理）─────────────────────────────
    def add_body(self, text_with_refs: str):
        """
        添加正文段落。

        自动识别文中的引用标记 [1] [2] [3][4] 等，将其转换为：
          - 上标样式（右上角小字）
          - 点击可跳转到参考文献对应条目（需配合 add_reference 使用）

        Parameters
        ----------
        text_with_refs : 包含 [n] 格式引用的段落文本。
        """
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.first_line_indent = Cm(0.74)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

        # 拆分文本和引用标记
        pattern = r'\[(\d+(?:[,\s]*\d+)*)\]'
        parts = re.split(pattern, text_with_refs)

        for i, part in enumerate(parts):
            if part == '':
                continue
            if i % 2 == 0:
                # 普通文本
                self._add_run(para, part)
            else:
                # 引用数字（可能逗号分隔多个，如 "3,4"）
                ref_nums = re.findall(r'\d+', part)
                for rn in ref_nums:
                    self._insert_citation_hyperlink(para, rn)

        return para

    def _insert_citation_hyperlink(self, para, ref_num: str):
        """向段落插入一个上标超链接引用标记 [N]。"""
        hl = etree.SubElement(para._element, _tag('hyperlink'))
        hl.set(_tag('anchor'), f'ref{ref_num}')
        hl.set(_tag('history'), '1')

        r = etree.SubElement(hl, _tag('r'))
        rPr = etree.SubElement(r, _tag('rPr'))

        va = etree.SubElement(rPr, _tag('vertAlign'))
        va.set(_tag('val'), 'superscript')

        rf = etree.SubElement(rPr, _tag('rFonts'))
        rf.set(_tag('eastAsia'), '宋体')
        rf.set(_tag('ascii'), '宋体')
        rf.set(_tag('hAnsi'), '宋体')

        sz = etree.SubElement(rPr, _tag('sz'))
        sz.set(_tag('val'), '21')

        t = etree.SubElement(r, _tag('t'))
        t.set(f'{{{XML_NS}}}space', 'preserve')
        t.text = f'[{ref_num}]'

    # ── 简单段落（无引用）────────────────────────────────
    def add_simple_para(self, text: str, bold: bool = False,
                        font_size: Pt = Pt(12),
                        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                        first_line_indent=Cm(0.74)):
        """添加不含引用的普通段落。"""
        para = self.doc.add_paragraph()
        para.alignment = alignment
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.first_line_indent = first_line_indent
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        self._add_run(para, text, font_size=font_size, bold=bold)
        return para

    # ── 参考文献 ──────────────────────────────────────────
    def add_reference_section_title(self):
        """添加「参考文献」标题（黑体小四居中），前面自动空一行。"""
        self._blank_line()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run('参考文献')
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(12)
        self._refs_added = True

    def add_reference(self, text: str, ref_num: int):
        """
        添加一条参考文献条目（宋体五号），并插入书签锚点供正文引用跳转。

        Parameters
        ----------
        text : 完整的参考文献条目文本（不含序号前缀，如 "作者.题名[J].刊名..."）
        ref_num : 参考文献序号（1, 2, 3...）
        """
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)

        bm_name = f'ref{ref_num}'
        bm_id = _next_bm_id()

        bm_start = etree.Element(_tag('bookmarkStart'))
        bm_start.set(_tag('id'), str(bm_id))
        bm_start.set(_tag('name'), bm_name)
        para._element.append(bm_start)

        run = para.add_run(f'[{ref_num}] {text}')
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10.5)

        bm_end = etree.Element(_tag('bookmarkEnd'))
        bm_end.set(_tag('id'), str(bm_id))
        para._element.append(bm_end)

        return para

    # ── 空白行 ────────────────────────────────────────────
    def add_blank(self):
        """添加一个空白段落（可用于章节间距调整）。"""
        self._blank_line()

    # ── 保存 ──────────────────────────────────────────────
    def save(self, filepath: str):
        """保存 Word 文档到指定路径。"""
        self.doc.save(filepath)
        print(f'[PaperGenerator] 论文已保存至: {filepath}')

    # ── 字数统计 ──────────────────────────────────────────
    def count_chars(self) -> dict:
        """返回 {chinese, english, total} 字数字典。"""
        full_text = []
        for para in self.doc.paragraphs:
            full_text.append(para.text)
        all_text = ''.join(full_text)
        chinese = len(re.findall(r'[一-鿿]', all_text))
        english = len(re.findall(r'[a-zA-Z]+', all_text))
        return {'chinese': chinese, 'english': english, 'total': chinese + english}


# ═══════════════════════════════════════════════════════════════
# 示例：用本脚本直接生成一篇演示论文
# ═══════════════════════════════════════════════════════════════

if __name__ == '__main__':
    gen = PaperGenerator()

    # 封面
    gen.set_cover_info(
        course='品牌文化研究与传播',
        school='艺术设计学院',
        student_id='M07xxxxxxx',
        author='xxx',
        advisor='陈红艳'
    )
    gen.set_title('沉浸式商业空间中的品牌体验设计与传播创新研究',
                  subtitle='——以北京SKP-S为例')

    # 内页
    gen.set_inner_header(
        '沉浸式商业空间中的品牌体验设计与传播创新研究——以北京SKP-S为例',
        'M07xxxxxxx xxx'
    )

    # 摘要
    gen.set_abstract(
        '随着体验经济时代的到来，沉浸式商业空间作为一种融合艺术、技术与消费的'
        '新型零售形态，正在重塑品牌与消费者之间的关系。本文以北京SKP-S为研究对象，'
        '基于沉浸式体验理论与品牌传播理论，从空间叙事、视觉系统、交互体验和策展式'
        '零售四个维度，系统分析其在品牌传播中的创新设计策略。研究发现，SKP-S通过'
        '数字化艺术装置、主题化空间叙事和多感官交互设计，构建了具有高度沉浸感的'
        '品牌体验场景，实现了从功能消费向意义消费的范式转换。',
        keywords='沉浸式体验；品牌传播；商业空间；体验设计；SKP-S'
    )

    # ── 正文 ──
    gen.add_heading('一、绪论')
    gen.add_heading('（一）研究背景与意义', level=2)
    gen.add_body(
        '在体验经济蓬勃发展的背景下，消费者的需求已从物质功能层面转向情感体验'
        '与意义建构层面。Pine与Gilmore在《体验经济》一书中指出，体验作为一种独立'
        '的经济提供物，正成为企业价值创造的核心载体[1]。零售商业空间作为品牌与'
        '消费者接触的关键触点，其功能定位正在发生根本性转变。曹智辉等人的研究'
        '进一步指出，沉浸式体验场景的建构已成为品牌竞争的新高地[2]。'
    )
    gen.add_body(
        '北京SKP-S作为中国沉浸式商业空间的标志性项目，自2019年开业以来便以其'
        '颠覆性的空间设计理念引发广泛关注。冯乐群与李维在研究中指出，SKP-S营造了'
        '以"数字-模拟未来"为主题的沉浸式科幻购物场景[3]。张悦群则从文化生产逻辑'
        '角度进行了批判性审视[4]。'
    )

    gen.add_heading('（二）研究对象与方法', level=2)
    gen.add_body(
        '本文以北京SKP-S为研究对象，采用案例研究法、文献分析法和实地观察法相结合'
        '的研究路径，从空间叙事设计、视觉识别系统、交互体验设计和策展式零售四个'
        '维度展开系统剖析。'
    )

    gen.add_heading('二、结论')
    gen.add_body(
        '研究表明，沉浸式商业空间通过整体空间叙事、数字美学视觉系统、多感官交互'
        '设计和策展式零售策略，实现了品牌传播的范式转换[3][4]。其核心启示在于：'
        '品牌传播的核心竞争力在于创造不可替代的深度体验场景[7][8]。'
    )

    # 参考文献
    gen.add_reference_section_title()
    gen.add_reference(
        'Pine B J, Gilmore J H. The Experience Economy[M]. Boston: '
        'Harvard Business School Press, 1999: 1-25.',
        ref_num=1
    )
    gen.add_reference(
        '曹智辉,妥艳媜,韩秋晨,等. 沉浸式体验场景的建构过程与机理[J]. '
        '外国经济与管理, 2024, 46(9): 48-65.',
        ref_num=2
    )
    gen.add_reference(
        '冯乐群,李维. 沉浸式体验与审美消费——以北京SKP-S商场中的当代艺术元素为例[J]. '
        '齐齐哈尔大学学报(哲学社会科学版), 2021(6): 122-126.',
        ref_num=3
    )
    gen.add_reference(
        '张悦群. 都市消费空间的"体验转向"及其文化生产逻辑[J]. '
        '中国图书评论, 2022(5): 82-91.',
        ref_num=4
    )
    gen.add_reference(
        '贺建平,黄秋皓. 消费者沉浸体验对品牌推崇的影响[J]. '
        '商业经济研究, 2023(6): 51-55.',
        ref_num=5
    )

    # 保存
    gen.save('demo_paper.docx')
    stats = gen.count_chars()
    print(f'字数统计: 中文 {stats["chinese"]} + 英文 {stats["english"]} '
          f'= 总计 {stats["total"]}')
